"""Tests unitarios de `tools.cv_tools` (helpers y `extract_candidate_data`)."""

import io
import json

import pytest

pytest.importorskip("boto3")
pytest.importorskip("docx")


@pytest.fixture(autouse=True)
def _s3_bucket_name_for_tests(monkeypatch):
    monkeypatch.setenv("S3_BUCKET_NAME", "test-bucket")


def test_get_s3_bucket_name_requires_env(monkeypatch):
    from tools import cv_tools

    monkeypatch.delenv("S3_BUCKET_NAME", raising=False)
    with pytest.raises(ValueError, match="S3_BUCKET_NAME"):
        cv_tools.get_s3_bucket_name()


def test_get_s3_client_raises_without_access_key(monkeypatch):
    from tools import cv_tools

    monkeypatch.delenv("AWS_ACCESS_KEY_ID", raising=False)
    monkeypatch.setenv("AWS_SECRET_ACCESS_KEY", "secret")
    with pytest.raises(ValueError, match="AWS_ACCESS_KEY_ID"):
        cv_tools._get_s3_client()


def test_get_s3_client_raises_without_secret(monkeypatch):
    from tools import cv_tools

    monkeypatch.setenv("AWS_ACCESS_KEY_ID", "AKIAKEY")
    monkeypatch.delenv("AWS_SECRET_ACCESS_KEY", raising=False)
    with pytest.raises(ValueError, match="AWS_SECRET_ACCESS_KEY"):
        cv_tools._get_s3_client()


def test_get_s3_client_returns_boto_client(monkeypatch):
    from tools import cv_tools

    monkeypatch.setenv("AWS_ACCESS_KEY_ID", "AKIATEST")
    monkeypatch.setenv("AWS_SECRET_ACCESS_KEY", "secretkey")
    monkeypatch.setenv("AWS_REGION", "us-east-1")

    class _FakeS3:
        pass

    def _fake_client(service, **_kwargs):
        assert service == "s3"
        return _FakeS3()

    monkeypatch.setattr(cv_tools.boto3, "client", _fake_client)
    client = cv_tools._get_s3_client()
    assert isinstance(client, _FakeS3)


def test_extract_text_from_docx_reads_paragraphs():
    from docx import Document

    from tools import cv_tools

    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Línea uno")
    doc.add_paragraph("Línea dos")
    doc.save(buf)
    text = cv_tools._extract_text_from_docx(buf.getvalue())
    assert "Línea uno" in text
    assert "Línea dos" in text


def test_extract_text_from_doc_delegates_to_docx():
    from docx import Document

    from tools import cv_tools

    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("DOC legacy body")
    doc.save(buf)
    text = cv_tools._extract_text_from_doc(buf.getvalue())
    assert "DOC legacy body" in text


def test_extract_candidate_data_finds_email_and_stack():
    from tools import cv_tools

    cv = """
    Jane Doe
    jane.doe@example.com
    Senior Python developer with Django experience.
    """
    out = json.loads(cv_tools.extract_candidate_data.func(cv))
    assert out["success"] is True
    hints = out["extracted_hints"]
    assert "jane.doe@example.com" in hints["emails_found"]
    assert "Python" in hints["technologies_found"]


def test_extract_candidate_data_error_returns_json(monkeypatch):
    from tools import cv_tools

    def _boom(*_a, **_k):
        raise RuntimeError("regex fail")

    monkeypatch.setattr(cv_tools.re, "findall", _boom)
    out = json.loads(cv_tools.extract_candidate_data.func("x"))
    assert out["success"] is False
    assert "regex fail" in out.get("error", "")


def _fake_s3_client():
    class _S3:
        class exceptions:
            NoSuchKey = type("NoSuchKey", (Exception,), {})
            NoSuchBucket = type("NoSuchBucket", (Exception,), {})

    return _S3()


def test_extract_text_from_pdf_uses_pdfplumber(monkeypatch):
    from tools import cv_tools

    class _Page:
        def extract_text(self):
            return "Texto desde pdfplumber"

    class _Pdf:
        pages = [_Page()]

        def __enter__(self):
            return self

        def __exit__(self, *a):
            return False

    monkeypatch.setattr(cv_tools.pdfplumber, "open", lambda _b: _Pdf())
    assert cv_tools._extract_text_from_pdf(b"%PDF-1.4 fake") == "Texto desde pdfplumber"


def test_extract_text_from_pdf_uses_pypdf2_when_plumber_empty(monkeypatch):
    from tools import cv_tools

    class _EmptyPage:
        def extract_text(self):
            return None

    class _EmptyPdf:
        pages = [_EmptyPage()]

        def __enter__(self):
            return self

        def __exit__(self, *a):
            return False

    class _PyPage:
        def extract_text(self):
            return "Texto PyPDF2"

    class _PyReader:
        pages = [_PyPage()]

    monkeypatch.setattr(cv_tools.pdfplumber, "open", lambda _b: _EmptyPdf())
    monkeypatch.setattr(cv_tools, "PdfReader", lambda _b: _PyReader())
    monkeypatch.setattr(cv_tools, "pdfminer_extract_text", lambda _b: "")
    monkeypatch.setattr(cv_tools, "_extract_text_from_pdf_with_textract", lambda _b: "")

    assert cv_tools._extract_text_from_pdf(b"%PDF-1.4 x") == "Texto PyPDF2"


def test_extract_text_from_pdf_uses_pdfminer_when_plumber_and_pypdf_empty(monkeypatch):
    from tools import cv_tools

    class _EmptyPage:
        def extract_text(self):
            return None

    class _EmptyPdf:
        pages = [_EmptyPage()]

        def __enter__(self):
            return self

        def __exit__(self, *a):
            return False

    class _PyReader:
        pages = [_EmptyPage()]

    monkeypatch.setattr(cv_tools.pdfplumber, "open", lambda _b: _EmptyPdf())
    monkeypatch.setattr(cv_tools, "PdfReader", lambda _b: _PyReader())
    monkeypatch.setattr(cv_tools, "pdfminer_extract_text", lambda _b: "  salida pdfminer  ")
    monkeypatch.setattr(cv_tools, "_extract_text_from_pdf_with_textract", lambda _b: "")

    assert cv_tools._extract_text_from_pdf(b"%PDF-1.4 y") == "salida pdfminer"


def test_extract_text_from_pdf_raises_when_all_methods_fail(monkeypatch):
    from tools import cv_tools

    class _EmptyPage:
        def extract_text(self):
            return None

    class _EmptyPdf:
        pages = [_EmptyPage()]

        def __enter__(self):
            return self

        def __exit__(self, *a):
            return False

    class _PyReader:
        pages = [_EmptyPage()]

    monkeypatch.setattr(cv_tools.pdfplumber, "open", lambda _b: _EmptyPdf())
    monkeypatch.setattr(cv_tools, "PdfReader", lambda _b: _PyReader())
    monkeypatch.setattr(cv_tools, "pdfminer_extract_text", lambda _b: "")
    monkeypatch.setattr(
        cv_tools,
        "_extract_text_from_pdf_with_textract",
        lambda _b: (_ for _ in ()).throw(RuntimeError("textract off")),
    )

    with pytest.raises(Exception, match="No se pudo extraer texto del PDF"):
        cv_tools._extract_text_from_pdf(b"%PDF-1.4 dead")


def test_extract_text_from_pdf_with_textract_access_denied(monkeypatch):
    from tools import cv_tools

    def _fake_textract(_service, **_kwargs):
        def _detect(**_kw):
            err = Exception("AccessDeniedException")
            raise err

        client = type("T", (), {"detect_document_text": _detect})()
        return client

    monkeypatch.setenv("AWS_ACCESS_KEY_ID", "k")
    monkeypatch.setenv("AWS_SECRET_ACCESS_KEY", "s")
    monkeypatch.setattr(cv_tools.boto3, "client", _fake_textract)

    with pytest.raises(Exception, match="Textract"):
        cv_tools._extract_text_from_pdf_with_textract(b"x")


def test_extract_text_from_pdf_with_textract_success(monkeypatch):
    from tools import cv_tools

    def _fake_client(service, **_kwargs):
        assert service == "textract"

        class _T:
            def detect_document_text(self, **_kw):
                return {
                    "Blocks": [
                        {"BlockType": "LINE", "Text": "Hola Textract"},
                        {"BlockType": "WORD", "Text": "x"},
                        {"BlockType": "LINE", "Text": "Segunda línea"},
                    ]
                }

        return _T()

    monkeypatch.setenv("AWS_ACCESS_KEY_ID", "k")
    monkeypatch.setenv("AWS_SECRET_ACCESS_KEY", "s")
    monkeypatch.setattr(cv_tools.boto3, "client", _fake_client)

    out = cv_tools._extract_text_from_pdf_with_textract(b"%PDF")
    assert "Hola Textract" in out
    assert "Segunda línea" in out


def test_extract_text_from_pdf_with_textract_invalid_parameter_message(monkeypatch):
    from tools import cv_tools

    def _fake_client(service, **_kwargs):
        class _T:
            def detect_document_text(self, **_kw):
                raise RuntimeError("InvalidParameterException: bad document")

        return _T()

    monkeypatch.setenv("AWS_ACCESS_KEY_ID", "k")
    monkeypatch.setenv("AWS_SECRET_ACCESS_KEY", "s")
    monkeypatch.setattr(cv_tools.boto3, "client", _fake_client)

    with pytest.raises(Exception, match="Parámetro inválido en Textract"):
        cv_tools._extract_text_from_pdf_with_textract(b"x")


def test_extract_text_from_pdf_with_textract_throughput_exceeded_message(monkeypatch):
    from tools import cv_tools

    def _fake_client(service, **_kwargs):
        class _T:
            def detect_document_text(self, **_kw):
                raise RuntimeError("ProvisionedThroughputExceededException: slow down")

        return _T()

    monkeypatch.setenv("AWS_ACCESS_KEY_ID", "k")
    monkeypatch.setenv("AWS_SECRET_ACCESS_KEY", "s")
    monkeypatch.setattr(cv_tools.boto3, "client", _fake_client)

    with pytest.raises(Exception, match="Límite de Textract"):
        cv_tools._extract_text_from_pdf_with_textract(b"x")


def test_extract_text_from_pdf_with_textract_no_line_blocks_raises(monkeypatch):
    from tools import cv_tools

    def _fake_client(service, **_kwargs):
        class _T:
            def detect_document_text(self, **_kw):
                return {
                    "Blocks": [
                        {"BlockType": "WORD", "Text": "solo"},
                    ]
                }

        return _T()

    monkeypatch.setenv("AWS_ACCESS_KEY_ID", "k")
    monkeypatch.setenv("AWS_SECRET_ACCESS_KEY", "s")
    monkeypatch.setattr(cv_tools.boto3, "client", _fake_client)

    with pytest.raises(Exception, match="Textract OCR"):
        cv_tools._extract_text_from_pdf_with_textract(b"x")


def test_extract_text_from_pdf_textract_failure_runs_pdfreader_diagnostic(monkeypatch):
    from tools import cv_tools

    class _EmptyPage:
        def extract_text(self):
            return None

    class _EmptyPdf:
        pages = [_EmptyPage()]

        def __enter__(self):
            return self

        def __exit__(self, *a):
            return False

    class _PyReader:
        pages = [_EmptyPage()]
        is_encrypted = False

    monkeypatch.setattr(cv_tools.pdfplumber, "open", lambda _b: _EmptyPdf())
    monkeypatch.setattr(cv_tools, "PdfReader", lambda _b: _PyReader())
    monkeypatch.setattr(cv_tools, "pdfminer_extract_text", lambda _b: "")
    monkeypatch.setattr(
        cv_tools,
        "_extract_text_from_pdf_with_textract",
        lambda _b: (_ for _ in ()).throw(RuntimeError("textract caído")),
    )

    with pytest.raises(Exception, match="No se pudo extraer texto del PDF"):
        cv_tools._extract_text_from_pdf(b"%PDF-1.4 diag")


def test_extract_text_from_pdf_textract_failure_diagnostic_pdfreader_raises(monkeypatch):
    from tools import cv_tools

    class _EmptyPage:
        def extract_text(self):
            return None

    class _EmptyPdf:
        pages = [_EmptyPage()]

        def __enter__(self):
            return self

        def __exit__(self, *a):
            return False

    class _PyReader:
        pages = [_EmptyPage()]

    _n = [0]

    def _reader(_b):
        _n[0] += 1
        if _n[0] == 1:
            return _PyReader()
        raise RuntimeError("diag open fail")

    monkeypatch.setattr(cv_tools.pdfplumber, "open", lambda _b: _EmptyPdf())
    monkeypatch.setattr(cv_tools, "PdfReader", _reader)
    monkeypatch.setattr(cv_tools, "pdfminer_extract_text", lambda _b: "")
    monkeypatch.setattr(
        cv_tools,
        "_extract_text_from_pdf_with_textract",
        lambda _b: (_ for _ in ()).throw(RuntimeError("textract off")),
    )

    with pytest.raises(Exception, match="No se pudo extraer texto del PDF"):
        cv_tools._extract_text_from_pdf(b"%PDF-1.4 diag2")


def test_extract_text_from_pdf_with_textract_file_too_large(monkeypatch):
    from tools import cv_tools

    big = b"x" * (6 * 1024 * 1024)
    with pytest.raises(Exception, match="muy grande"):
        cv_tools._extract_text_from_pdf_with_textract(big)


def test_download_cv_from_s3_docx_success(monkeypatch):
    from docx import Document

    from tools import cv_tools

    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Contenido CV")
    doc.save(buf)
    raw = buf.getvalue()

    class _Body:
        def read(self):
            return raw

    s3 = _fake_s3_client()
    s3.head_object = lambda **kwargs: {}
    s3.get_object = lambda **kwargs: {"Body": _Body()}

    monkeypatch.setattr(cv_tools, "_get_s3_client", lambda: s3)

    out = json.loads(cv_tools.download_cv_from_s3.func("mi_cv.docx"))
    assert out["success"] is True
    assert "Contenido CV" in out["text_content"]
    assert out["file_type"] == "docx"


def test_download_cv_from_s3_pdf_success(monkeypatch):
    from tools import cv_tools

    class _Body:
        def read(self):
            return b"%PDF-1.4"

    s3 = _fake_s3_client()
    s3.head_object = lambda **kwargs: {}
    s3.get_object = lambda **kwargs: {"Body": _Body()}

    monkeypatch.setattr(cv_tools, "_get_s3_client", lambda: s3)
    monkeypatch.setattr(cv_tools, "_extract_text_from_pdf", lambda _b: "PDF extraído")

    out = json.loads(cv_tools.download_cv_from_s3.func("cv.pdf"))
    assert out["success"] is True
    assert out["text_content"] == "PDF extraído"


def test_download_cv_from_s3_prefix_in_filename(monkeypatch):
    from docx import Document

    from tools import cv_tools

    buf = io.BytesIO()
    Document().save(buf)
    raw = buf.getvalue()

    class _Body:
        def read(self):
            return raw

    s3 = _fake_s3_client()
    s3.head_object = lambda **kwargs: {}
    s3.get_object = lambda **kwargs: {"Body": _Body()}

    monkeypatch.setattr(cv_tools, "_get_s3_client", lambda: s3)

    out = json.loads(cv_tools.download_cv_from_s3.func("cvs/pref.docx"))
    assert out["success"] is True
    assert "cvs/" in out["s3_key"]


def test_download_cv_from_s3_no_such_key_returns_json(monkeypatch):
    from tools import cv_tools

    s3 = _fake_s3_client()

    def _head(**_kwargs):
        raise s3.exceptions.NoSuchKey()

    s3.head_object = _head
    monkeypatch.setattr(cv_tools, "_get_s3_client", lambda: s3)

    out = json.loads(cv_tools.download_cv_from_s3.func("missing.docx"))
    assert out["success"] is False
    assert out["error_type"] == "FileNotFoundError"


def test_download_cv_from_s3_empty_file_returns_json(monkeypatch):
    from tools import cv_tools

    class _Body:
        def read(self):
            return b""

    s3 = _fake_s3_client()
    s3.head_object = lambda **kwargs: {}
    s3.get_object = lambda **kwargs: {"Body": _Body()}

    monkeypatch.setattr(cv_tools, "_get_s3_client", lambda: s3)

    out = json.loads(cv_tools.download_cv_from_s3.func("empty.docx"))
    assert out["success"] is False
    assert out["error_type"] == "ValueError"


def test_download_cv_from_s3_unsupported_format_returns_json(monkeypatch):
    from tools import cv_tools

    class _Body:
        def read(self):
            return b"data"

    s3 = _fake_s3_client()
    s3.head_object = lambda **kwargs: {}
    s3.get_object = lambda **kwargs: {"Body": _Body()}

    monkeypatch.setattr(cv_tools, "_get_s3_client", lambda: s3)

    out = json.loads(cv_tools.download_cv_from_s3.func("file.txt"))
    assert out["success"] is False
    assert "Unsupported file format" in out["error"]


def test_extract_candidate_data_suggested_profile_fullstack():
    from tools import cv_tools

    cv = """
    Full stack developer.
    React experience and Python Django backend.
    """
    out = json.loads(cv_tools.extract_candidate_data.func(cv))
    assert out["success"] is True
    assert out["extracted_hints"]["suggested_profile"] == "Fullstack"


def test_extract_candidate_data_role_quality_assurance():
    from tools import cv_tools

    cv = "quality assurance engineer at ACME Corp"
    out = json.loads(cv_tools.extract_candidate_data.func(cv))
    assert out["success"] is True
    hints = out["extracted_hints"]
    assert hints["suggested_profile"] == "QA"
    assert any("quality" in t.lower() for t in hints["role_titles_found"])


def test_extract_candidate_data_linkedin_normalized():
    from tools import cv_tools

    cv = "Perfil: linkedin.com/in/janedoe"
    out = json.loads(cv_tools.extract_candidate_data.func(cv))
    assert out["success"] is True
    urls = out["extracted_hints"]["linkedin_urls_found"]
    assert any(u.startswith("https://") for u in urls)


def test_download_cv_from_s3_no_such_bucket_on_get_object(monkeypatch):
    from tools import cv_tools

    s3 = _fake_s3_client()

    def _get(**_kwargs):
        raise s3.exceptions.NoSuchBucket()

    s3.head_object = lambda **kwargs: {}
    s3.get_object = _get
    monkeypatch.setattr(cv_tools, "_get_s3_client", lambda: s3)

    out = json.loads(cv_tools.download_cv_from_s3.func("x.docx"))
    assert out["success"] is False
    assert out["error_type"] == "NoSuchBucket"


def test_download_cv_from_s3_access_denied_message(monkeypatch):
    from tools import cv_tools

    class _Body:
        def read(self):
            return b"x"

    s3 = _fake_s3_client()
    s3.head_object = lambda **kwargs: {}

    def _get(**_kwargs):
        raise RuntimeError("AccessDenied: user not allowed")

    s3.get_object = _get
    monkeypatch.setattr(cv_tools, "_get_s3_client", lambda: s3)

    out = json.loads(cv_tools.download_cv_from_s3.func("a.docx"))
    assert out["success"] is False
    assert "Access Denied" in out["error"]


def test_download_cv_from_s3_403_message(monkeypatch):
    from tools import cv_tools

    s3 = _fake_s3_client()
    s3.head_object = lambda **kwargs: {}
    s3.get_object = lambda **kwargs: (_ for _ in ()).throw(RuntimeError("HTTP 403 Forbidden"))
    monkeypatch.setattr(cv_tools, "_get_s3_client", lambda: s3)

    out = json.loads(cv_tools.download_cv_from_s3.func("a.docx"))
    assert out["success"] is False
    assert "403" in out["error"] or "Access Denied" in out["error"]


def test_download_cv_from_s3_invalid_access_key_message(monkeypatch):
    from tools import cv_tools

    s3 = _fake_s3_client()
    s3.head_object = lambda **kwargs: {}
    s3.get_object = lambda **kwargs: (_ for _ in ()).throw(
        RuntimeError("InvalidAccessKeyId: bad key"),
    )
    monkeypatch.setattr(cv_tools, "_get_s3_client", lambda: s3)

    out = json.loads(cv_tools.download_cv_from_s3.func("a.docx"))
    assert out["success"] is False
    assert "AWS_ACCESS_KEY_ID" in out["error"]


def test_download_cv_from_s3_signature_mismatch_message(monkeypatch):
    from tools import cv_tools

    s3 = _fake_s3_client()
    s3.head_object = lambda **kwargs: {}
    s3.get_object = lambda **kwargs: (_ for _ in ()).throw(
        RuntimeError("SignatureDoesNotMatch"),
    )
    monkeypatch.setattr(cv_tools, "_get_s3_client", lambda: s3)

    out = json.loads(cv_tools.download_cv_from_s3.func("a.docx"))
    assert out["success"] is False
    assert "AWS_SECRET_ACCESS_KEY" in out["error"]


def test_download_cv_from_s3_no_such_key_in_error_message(monkeypatch):
    from tools import cv_tools

    s3 = _fake_s3_client()
    s3.head_object = lambda **kwargs: {}
    s3.get_object = lambda **kwargs: (_ for _ in ()).throw(
        RuntimeError("NoSuchKey: object gone during read"),
    )
    monkeypatch.setattr(cv_tools, "_get_s3_client", lambda: s3)

    out = json.loads(cv_tools.download_cv_from_s3.func("gone.docx"))
    assert out["success"] is False
    assert "s3://" in out["error"]
    assert "gone.docx" in out["error"] or "cvs/gone" in out["error"]


def test_download_cv_from_s3_generic_exception_message(monkeypatch):
    from tools import cv_tools

    s3 = _fake_s3_client()
    s3.head_object = lambda **kwargs: {}
    s3.get_object = lambda **kwargs: (_ for _ in ()).throw(RuntimeError("network glitch"))
    monkeypatch.setattr(cv_tools, "_get_s3_client", lambda: s3)

    out = json.loads(cv_tools.download_cv_from_s3.func("a.docx"))
    assert out["success"] is False
    assert out["error"] == "network glitch"


def test_download_cv_from_s3_head_object_wraps_exception(monkeypatch):
    from tools import cv_tools

    s3 = _fake_s3_client()
    s3.head_object = lambda **kwargs: (_ for _ in ()).throw(RuntimeError("throttle"))
    monkeypatch.setattr(cv_tools, "_get_s3_client", lambda: s3)

    out = json.loads(cv_tools.download_cv_from_s3.func("a.docx"))
    assert out["success"] is False
    assert "Error verificando objeto en S3" in out["error"]


def test_extract_candidate_data_suggested_profile_ux_ui():
    from tools import cv_tools

    cv = "Product designer, Figma expert, UX wireframes"
    out = json.loads(cv_tools.extract_candidate_data.func(cv))
    assert out["success"] is True
    assert out["extracted_hints"]["suggested_profile"] == "UX/UI"


def test_extract_candidate_data_suggested_profile_devops():
    from tools import cv_tools

    cv = "AWS solutions architect, Docker, Kubernetes, Terraform pipelines"
    out = json.loads(cv_tools.extract_candidate_data.func(cv))
    assert out["success"] is True
    assert out["extracted_hints"]["suggested_profile"] == "DevOps"


def test_extract_candidate_data_suggested_profile_team_manager_business():
    from tools import cv_tools

    cv = "Analista contable, Excel avanzado, Power BI, reportes para dirección"
    out = json.loads(cv_tools.extract_candidate_data.func(cv))
    assert out["success"] is True
    assert out["extracted_hints"]["suggested_profile"] == "Team Manager"


def test_extract_candidate_data_suggested_profile_frontend_only():
    from tools import cv_tools

    cv = "Senior React developer, Vue components, CSS modules"
    out = json.loads(cv_tools.extract_candidate_data.func(cv))
    assert out["success"] is True
    assert out["extracted_hints"]["suggested_profile"] == "Frontend"


def test_extract_candidate_data_frontend_not_fullstack_with_javascript_stack():
    """Regresión: JavaScript + React no debe implicar backend por subcadena 'java'."""
    from tools import cv_tools

    cv = "Senior React developer, TypeScript, Next.js SPA, JavaScript ES2020"
    out = json.loads(cv_tools.extract_candidate_data.func(cv))
    assert out["success"] is True
    assert out["extracted_hints"]["suggested_profile"] == "Frontend"


def test_extract_candidate_data_fullstack_requires_real_backend_signal():
    from tools import cv_tools

    cv = "Full stack: React frontend, Python and Django APIs, PostgreSQL"
    out = json.loads(cv_tools.extract_candidate_data.func(cv))
    assert out["success"] is True
    assert out["extracted_hints"]["suggested_profile"] == "Fullstack"


def test_extract_candidate_data_suggested_profile_backend_only():
    from tools import cv_tools

    cv = "Backend engineer, Python, Django REST, PostgreSQL"
    out = json.loads(cv_tools.extract_candidate_data.func(cv))
    assert out["success"] is True
    assert out["extracted_hints"]["suggested_profile"] == "Backend"


def test_extract_candidate_data_role_recursos_humanos_team_manager():
    from tools import cv_tools

    cv = "Coordinadora de recursos humanos, onboarding y clima laboral"
    out = json.loads(cv_tools.extract_candidate_data.func(cv))
    assert out["success"] is True
    assert out["extracted_hints"]["suggested_profile"] == "Team Manager"


def test_extract_candidate_data_suggested_profile_team_manager_tech_lead():
    from tools import cv_tools

    cv = "tech lead de ingeniería, Scrum, liderazgo de equipo"
    out = json.loads(cv_tools.extract_candidate_data.func(cv))
    assert out["success"] is True
    assert out["extracted_hints"]["suggested_profile"] == "Team Manager"
